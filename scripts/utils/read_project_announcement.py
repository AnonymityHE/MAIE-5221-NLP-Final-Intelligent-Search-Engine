"""
读取Project Announcement.docx文档内容
"""
import sys
from pathlib import Path

# 添加项目根目录到Python路径
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

try:
    from docx import Document
    
    doc_path = "docs/Project Announcement.docx"
    doc = Document(doc_path)
    
    print("=" * 80)
    print("Project Announcement 内容:")
    print("=" * 80)
    print()
    
    for i, para in enumerate(doc.paragraphs, 1):
        text = para.text.strip()
        if text:  # 只打印非空段落
            print(f"[段落 {i}]")
            print(text)
            print()
    
    # 也读取表格内容
    if doc.tables:
        print("=" * 80)
        print("表格内容:")
        print("=" * 80)
        for table_idx, table in enumerate(doc.tables, 1):
            print(f"\n表格 {table_idx}:")
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                print(" | ".join(row_text))
    
except ImportError:
    print("错误: 未安装 python-docx")
    print("请运行: pip install python-docx")
except FileNotFoundError:
    print(f"错误: 找不到文件 {doc_path}")
except Exception as e:
    print(f"读取文档时出错: {e}")

